"""
file_reader.py
Extracts text from any file type given raw bytes + filename.
Supported: PDF, DOCX, DOC, images (OCR), TXT, MSG
"""

import io
import logging
import tempfile
import os

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]

    try:
        if ext == "pdf":
            return _read_pdf(file_bytes)
        elif ext in ("docx",):
            return _read_docx(file_bytes)
        elif ext in ("doc",):
            return _read_doc(file_bytes)
        elif ext in ("txt", "csv", "eml"):
            return file_bytes.decode("utf-8", errors="replace")
        elif ext in ("png", "jpg", "jpeg", "tiff", "tif", "bmp", "gif"):
            return _ocr_image(file_bytes)
        elif ext in ("xlsx", "xls"):
            return _read_excel(file_bytes, ext)
        else:
            logger.warning(f"Unsupported file type: {ext} ({filename})")
            return f"[Could not extract text from {filename} — unsupported type: {ext}]"
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}")
        return f"[Error reading {filename}: {e}]"


def _read_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            pages.append(f"[Page {i+1}]\n{text}")
        else:
            # Scanned — OCR via image
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            ocr_text = _ocr_image(img_bytes)
            if ocr_text.strip():
                pages.append(f"[Page {i+1} (OCR)]\n{ocr_text}")
    return "\n\n".join(pages)


def _read_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _read_doc(data: bytes) -> str:
    """Convert legacy .doc via LibreOffice, then read as docx."""
    import subprocess
    with tempfile.TemporaryDirectory() as tmpdir:
        src = os.path.join(tmpdir, "file.doc")
        with open(src, "wb") as f:
            f.write(data)
        try:
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx", "--outdir", tmpdir, src],
                timeout=30, check=True, capture_output=True
            )
            out = os.path.join(tmpdir, "file.docx")
            if os.path.exists(out):
                with open(out, "rb") as f:
                    return _read_docx(f.read())
        except Exception as e:
            logger.warning(f"LibreOffice conversion failed: {e}")
    return "[Could not read .doc file]"


def _ocr_image(data: bytes) -> str:
    import pytesseract
    from PIL import Image
    img = Image.open(io.BytesIO(data))
    return pytesseract.image_to_string(img)


def _read_excel(data: bytes, ext: str) -> str:
    import pandas as pd
    engine = "openpyxl" if ext == "xlsx" else "xlrd"
    xf = pd.ExcelFile(io.BytesIO(data), engine=engine)
    parts = []
    for sheet in xf.sheet_names:
        df = xf.parse(sheet)
        parts.append(f"[Sheet: {sheet}]\n{df.to_string(index=False)}")
    return "\n\n".join(parts)
