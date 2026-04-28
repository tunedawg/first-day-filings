from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


FONT_NAME = "Century Schoolbook"
FONT_SIZE_PT = 13


REPLACE_EXACT = {
    "4600 Madison Ave. Ste. 810": "{{firmAddressBlock}}",
    "Kansas City, Missouri 64112": "",
    "(816) 809-2100": "",
    "ee@keenanfirm.com": "{{attorneyEmailBlock}}",
    "hilary@keenanfirm.com": "",
    "Attorneys for Plaintiff Fern Payne": "{{attorneyForLine}}",
    "/s/ Aaron Hadlow": "/s/{{signingAttorney}}",
    "An Attorney for Plaintiff {{plaintiffName}}": "{{attorneyForLine}}",
    "Attorney for Plaintiff {{plaintiffName}}": "{{attorneyForLine}}",
    "An Attorney for Plaintiff {{plaintiffName}}": "{{attorneyForLine}}",
}

REPLACE_SUBSTRINGS = {
    "THE METROPOLITAN\nCOMMUNITY COLLEGE\nFOUNDATION et al.,": "{{defendantName}}",
    "THE METROPOLITAN\nCOMMUNITY COLLEGE\nFOUNDATION, et al.,": "{{defendantName}}",
    "THE METROPOLITAN COMMUNITY COLLEGE FOUNDATION et al.": "{{defendantName}}",
    "THE METROPOLITAN COMMUNITY COLLEGE FOUNDATION, et al.": "{{defendantName}}",
    "Case No. \n2616-CV12184": "Case No. \n{{caseNumber}}",
    "Case No.: 2616-CV12184": "Case No.: {{caseNumber}}",
    "Case No. 2616-CV12184": "Case No. {{caseNumber}}",
    "Fern Payne": "{{plaintiffName}}",
    "Ms. Payne": "{{plaintiffReferenceName}}",
}

TABLE_CELL_REPLACEMENTS = {
    "THE METROPOLITAN\nCOMMUNITY COLLEGE\nFOUNDATION et al.,": "{{defendantName}}",
    "THE METROPOLITAN\nCOMMUNITY COLLEGE\nFOUNDATION, et al.,": "{{defendantName}}",
    "Case No. \n2616-CV12184": "Case No. \n{{caseNumber}}",
}

DEDUPABLE_PLACEHOLDERS = {
    "{{firmAddressBlock}}",
    "{{attorneyEmailBlock}}",
    "{{attorneyForLine}}",
    "{{attorneyRosterBlock}}",
}

FILE_SPECIFIC_PARAGRAPH_RULES = {
    "interrogatories-template.docx": [
        (
            "Please identify all complaints, whether formal or informal, made by any employee against Ms. Moore-Jones",
            "{{interrogatoryActorComplaintParagraphsBlock}}",
        ),
        (
            "With respect to the TRIO Program’s annual key metrics reports submitted to the United States Department of Education",
            "{{interrogatoryTrioMetricsParagraph}}",
        ),
    ],
    "rfps-template.docx": [
        (
            "Please produce all communications between Ms. Payne and Ms. Moore-Jones",
            "{{rfpPlaintiffCommunicationsBlock}}",
        ),
        (
            "Please produce all communications between {{plaintiffReferenceName}} and Ms. Moore-Jones",
            "{{rfpPlaintiffCommunicationsBlock}}",
        ),
        (
            "Please produce all communications between Ms. Payne and Ms. Jones Scott",
            "",
        ),
        (
            "Please produce all communications between {{plaintiffReferenceName}} and Ms. Jones Scott",
            "",
        ),
        (
            "Please produce all communications between Ms. Payne and Ms. Fockler",
            "",
        ),
        (
            "Please produce all communications between {{plaintiffReferenceName}} and Ms. Fockler",
            "",
        ),
        (
            "Please produce all communications between Ms. Payne and Ms. Hafner department or any employee therein",
            "",
        ),
        (
            "Please produce all communications between {{plaintiffReferenceName}} and Ms. Hafner department or any employee therein",
            "",
        ),
        (
            "Please produce all communications between Ms. Payne and Ms. Bradley",
            "",
        ),
        (
            "Please produce all communications between {{plaintiffReferenceName}} and Ms. Bradley",
            "",
        ),
        (
            "Please produce all communications between Ms. Payne and Mr. Wise",
            "",
        ),
        (
            "Please produce all communications between {{plaintiffReferenceName}} and Mr. Wise",
            "",
        ),
        (
            "Please produce all communications between Ms. Payne and Ms. Perez",
            "",
        ),
        (
            "Please produce all communications between {{plaintiffReferenceName}} and Ms. Perez",
            "",
        ),
        (
            "Please produce all communications between Ms. Payne and MCC’s Human Resources department or any employee therein",
            "",
        ),
        (
            "Please produce all communications between {{plaintiffReferenceName}} and MCC’s Human Resources department or any employee therein",
            "",
        ),
        (
            "Please produce all communications and documents which relate to any complaints made against or about Ms. Moore-Jones",
            "{{rfpActorComplaintParagraphsBlock}}",
        ),
        (
            "Please produce all communications and documents which relate to any complaints made against or about Dr. Jones Scott",
            "",
        ),
        (
            "Please produce all documents that reflect or relate to any coaching, performance counseling, or other performance-related guidance, whether formal or informal, which has been given to anyone under Ms. Moore-Jones’s supervision",
            "{{rfpSupervisorCoachingParagraph}}",
        ),
        (
            "Please produce all documents which constitute lists or charts of all employees supervised or overseen in any way by Ms. Moore-Jones",
            "{{rfpSupervisorEmployeeListParagraph}}",
        ),
        (
            "Please produce all documents which reflect the circumstances of the separation from employment of any employee who reported to Ms. Moore-Jones",
            "{{rfpSupervisorSeparationParagraph}}",
        ),
        (
            "Please produce all documents and communication related to placing Ms. Payne on a Formal Verbal Warning",
            "{{rfpPlaintiffDisciplineParagraphsBlock}}",
        ),
        (
            "Please produce all documents and communication related to placing {{plaintiffReferenceName}} on a Formal Verbal Warning",
            "{{rfpPlaintiffDisciplineParagraphsBlock}}",
        ),
        (
            "Please produce all documents and communication related to placing Ms. Payne on a Performance Improvement Plan",
            "",
        ),
        (
            "Please produce all documents and communication related to placing {{plaintiffReferenceName}} on a Performance Improvement Plan",
            "",
        ),
        (
            "Please produce all documents reflecting training that Ms. Moore-Jones completed",
            "{{rfpSupervisorTrainingParagraph}}",
        ),
        (
            "Please produce all documents reflecting training that any other person with Ms. Moore-Jones’s job title completed",
            "{{rfpSupervisorPeerTrainingParagraph}}",
        ),
        (
            "Produce all annual performance reports, key metrics reports, and any other reports submitted by Metropolitan Community College",
            "{{rfpTrioIssueParagraphsBlock}}",
        ),
        (
            "Produce all documents, spreadsheets, databases, enrollment rosters, and data compilations used to generate, calculate, or verify the student enrollment numbers",
            "",
        ),
        (
            "Produce all communications between MCC and the United States Department of Education",
            "",
        ),
        (
            "Produce all communications, including emails, text messages, instant messages, voicemails, and written notes, between Gabrielle Moore-Jones",
            "",
        ),
        (
            "Produce all documents related to MCC’s internal investigation conducted following Ms. Payne’s August 2023 report",
            "",
        ),
        (
            "Produce all documents related to MCC’s internal investigation conducted following {{plaintiffReferenceName}}’s August 2023 report",
            "",
        ),
        (
            "Produce all documents related to MCC’s internal investigation conducted following Ms. Payne’s August 2024 report",
            "",
        ),
        (
            "Produce all documents related to MCC’s internal investigation conducted following {{plaintiffReferenceName}}’s August 2024 report",
            "",
        ),
        (
            "Please produce all documents that reflect the actual funding received from the federal government to the TRIO program",
            "",
        ),
        (
            "Produce all training materials, policy documents, procedures, guidelines, or instructions provided to TRIO staff",
            "",
        ),
        (
            "Produce all documents, communications, or records reflecting any instruction, direction, or request made by Gabrielle Moore-Jones",
            "",
        ),
        (
            "Produce all documents reflecting any audit, review, inquiry, correction, amendment, or resubmission of any TRIO key metrics report",
            "",
        ),
    ],
}


def replace_in_paragraph(paragraph) -> bool:
    original = paragraph.text
    updated = original

    if updated in REPLACE_EXACT:
      updated = REPLACE_EXACT[updated]

    for needle, replacement in REPLACE_SUBSTRINGS.items():
        updated = updated.replace(needle, replacement)

    if updated == original:
        return False

    # Simpler than run-preserving surgery and good enough for template drafts.
    paragraph.text = updated
    return True


def apply_file_specific_rules(paragraph, rules) -> bool:
    original = paragraph.text
    updated = original

    for prefix, replacement in rules:
        if updated.strip().startswith(prefix):
            updated = replacement
            break

    if updated == original:
        return False

    paragraph.text = updated
    return True


def normalize_cell_text(text: str) -> str:
    updated = text
    for needle, replacement in TABLE_CELL_REPLACEMENTS.items():
        updated = updated.replace(needle, replacement)
    return updated


def dedupe_placeholder_paragraphs(paragraphs) -> int:
    changes = 0
    seen_recent = []

    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if text in DEDUPABLE_PLACEHOLDERS and text in seen_recent:
            paragraph.text = ""
            changes += 1
            continue

        seen_recent.append(text)
        seen_recent = seen_recent[-4:]

    return changes


def set_run_font(run) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.rFonts
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.append(run_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        run_fonts.set(qn(f"w:{attr}"), FONT_NAME)


def style_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        set_run_font(run)


def style_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph)
            for nested in cell.tables:
                style_table(nested)


def style_document(document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style.font.size = Pt(FONT_SIZE_PT)

    for paragraph in document.paragraphs:
        style_paragraph(paragraph)

    for table in document.tables:
        style_table(table)


def clear_paragraph(paragraph) -> None:
    paragraph.text = ""


def append_field_run(paragraph, field_code: str):
    run = paragraph.add_run()
    set_run_font(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_code

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)
    return run


def apply_footer_page_numbers(document) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        append_field_run(paragraph, "PAGE")
        style_paragraph(paragraph)

        for extra_paragraph in footer.paragraphs[1:]:
            clear_paragraph(extra_paragraph)
            style_paragraph(extra_paragraph)

        for table in footer.tables:
            style_table(table)


def replace_in_table(table) -> int:
    changes = 0
    for row in table.rows:
        for cell in row.cells:
            new_cell_text = normalize_cell_text(cell.text)
            if new_cell_text != cell.text:
                cell.text = new_cell_text
                changes += 1
            for paragraph in cell.paragraphs:
                if replace_in_paragraph(paragraph):
                    changes += 1
            changes += dedupe_placeholder_paragraphs(cell.paragraphs)
            for nested in cell.tables:
                changes += replace_in_table(nested)
    return changes


def normalize_docx(path: Path, output_dir: Path) -> int:
    document = Document(path)
    changes = 0
    rules = FILE_SPECIFIC_PARAGRAPH_RULES.get(path.name, [])

    for paragraph in document.paragraphs:
        if replace_in_paragraph(paragraph):
            changes += 1
        if rules and apply_file_specific_rules(paragraph, rules):
            changes += 1

    changes += dedupe_placeholder_paragraphs(document.paragraphs)

    for table in document.tables:
        changes += replace_in_table(table)

    style_document(document)
    apply_footer_page_numbers(document)

    output_dir.mkdir(parents=True, exist_ok=True)
    document.save(output_dir / path.name)
    return changes


def main() -> None:
    base_dir = Path(__file__).resolve().parents[1]
    input_dir = base_dir / "templates" / "docx-masters"
    output_dir = base_dir / "templates" / "final-docx-masters"

    for path in sorted(input_dir.glob("*.docx")):
        changes = normalize_docx(path, output_dir)
        print(f"{path.name}: {changes} changes")


if __name__ == "__main__":
    main()
