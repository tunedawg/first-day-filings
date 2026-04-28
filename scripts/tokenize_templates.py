"""
tokenize_templates.py

Reads the 6 Payne DOCX files from templates/final-docx-masters-v2/, applies
token replacements, and saves the results to templates/final-docx-masters-v2-tokenized/
with the standard filenames that upload_templates.js expects.

Run:
    python scripts/tokenize_templates.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
INPUT_DIR = ROOT / "templates" / "final-docx-masters-v2"
OUTPUT_DIR = ROOT / "templates" / "final-docx-masters-v2-tokenized"

# Maps the Payne-specific filenames to the standard template names that
# upload_templates.js and registry.json use.
FILE_MAP = {
    "Payne_MCC- Omnibus Notice of Deposition.docx": "omnibus-notice-template.docx",
    "Payne_MCC - Notice of Corporate Representative Deposition.docx": "corporate-representative-template.docx",
    "Payne_MCC - P's 1st RFPs to MCC (1).docx": "rfps-template.docx",
    "Payne_MCC - P's 1st ROGs to MCC (1).docx": "interrogatories-template.docx",
    "Payne_MCC - Proposed Protective Order.docx": "protective-order-template.docx",
    "Payne_MCC - Motion for Entry of Protective Order (1).docx": "motion-for-protective-order-template.docx",
}

# -------------------------------------------------------------------
# Global substring replacements applied to every paragraph in every doc.
# Order matters: more specific patterns first.
# -------------------------------------------------------------------
SUBSTRINGS = [
    # Court caption (some docs omit ", MISSOURI" on this line)
    ("IN THE CIRCUIT COURT OF JACKSON COUNTY, MISSOURI", "{{courtName}}"),
    ("IN THE CIRCUIT COURT OF JACKSON COUNTY", "{{courtName}}"),
    ("AT KANSAS CITY", "{{courtDivision}}"),
    # Parties / case
    ("Fern Payne", "{{plaintiffName}}"),
    ("Ms. Payne", "{{plaintiffReferenceName}}"),
    ("2616-CV12184", "{{caseNumber}}"),
    ("April 20, 2026", "{{serviceDate}}"),
    # Defendant long-form (appears in protective order and RFP intro)
    (
        "The Metropolitan Community College Foundation, Metropolitan Community College,"
        " and The Junior College District of Metropolitan Kansas City",
        "{{allDefendants}}",
    ),
    (
        "Metropolitan Community College Foundation, Metropolitan Community College,"
        " and The Junior College District of Metropolitan Kansas City",
        "{{allDefendants}}",
    ),
    # Defendant short-form reference used throughout body text
    ("“MCC” or “Defendants”", '"{{defendantReferenceName}}" or "Defendants"'),
    # Protective order judge
    ("HON. ADAM CAINE", "{{judgeName}}"),
    # Protective order action name
    (
        "Fern Payne V. The Metropolitan Community College Foundation, Case No. {{caseNumber}}",
        "{{protectiveOrderActionName}}",
    ),
    # Video operator (omnibus + corp rep)
    (
        "Noah Tunis or another authorized employee of the law firm of Keenan & Bhatia, LLC",
        "{{videoOperatorName}} or another authorized employee of the law firm of Keenan & Bhatia, LLC",
    ),
    ("Noah Tunis of Keenan & Bhatia", "{{videoOperatorName}} of Keenan & Bhatia"),
    ("Noah Tunis or an authorized representative of Keenan & Bhatia, LLC", "{{videoOperatorName}} or an authorized representative of Keenan & Bhatia, LLC"),
    # Firm address block (paragraph 1 of 3 — the other two are blanked below)
    ("4600 Madison Ave. Ste. 810", "{{firmAddressBlock}}"),
    ("Kansas City, Missouri 64112", ""),
    ("(816) 809-2100", ""),
    # Signing attorney
    ("/s/ Aaron Hadlow", "/s/{{signingAttorney}}"),
    ("/s/Aaron Hadlow", "/s/{{signingAttorney}}"),
    # Attorney-for line (must come after Fern Payne → {{plaintiffName}})
    ("Attorneys for Plaintiff {{plaintiffName}}", "{{attorneyForLine}}"),
    ("Attorneys for {{plaintiffName}}", "{{attorneyForLine}}"),
    ("Attorney for Plaintiff {{plaintiffName}}", "{{attorneyForLine}}"),
    ("An Attorney for Plaintiff {{plaintiffName}}", "{{attorneyForLine}}"),
    ("Attorney for Plaintiff", "{{attorneyForLine}}"),
    ("An Attorney for {{plaintiffName}}", "{{attorneyForLine}}"),
    # Catch any remaining plain references before tokenization made them generic
    ("Attorneys for Plaintiff Fern Payne", "{{attorneyForLine}}"),
    ("Attorney for Plaintiff Fern Payne", "{{attorneyForLine}}"),
    # Email block (first email becomes the token; subsequent emails are blanked)
    ("sonal@keenanfirm.com", "{{attorneyEmailBlock}}"),
    ("ee@keenanfirm.com", ""),
    ("jr@keenanfirm.com", ""),
    ("aaron@keenanfirm.com", ""),
    ("hilary@keenanfirm.com", ""),
    # Corp-rep contact (appears before email block replacement)
    ("E.E. Keenan at ee@keenanfirm.com", "{{zoomContactName}} at {{zoomContactEmail}}"),
    # Zoom contact format used in corp rep date/time paragraph
    ("Please contact E.E. Keenan at", "Please contact {{zoomContactName}} at"),
    # Corp rep entity
    ("Corporate Representative(s) for Defendants", "{{corpRepEntity}}"),
    # Corp rep date/time
    (
        "As agreed by the parties or in the absence of such agreement, July 2, 2026, at 10:00 am Central Time.",
        "{{corpRepDateTime}}",
    ),
    ("Known to Defendants.", "{{corpRepLocation}}"),
]

# Exact-match whole-paragraph replacements (checked after SUBSTRINGS).
EXACT = {
    # Attorney roster lines — first becomes the block token, others blank.
    "Edward (E.E.) Keenan \t(Mo. #62993)": "{{attorneyRosterBlock}}",
    "Edward (E.E.) Keenan  (Mo. #62993)": "{{attorneyRosterBlock}}",
    "Sonal Bhatia\t(Mo. #67519)": "",
    "Sonal Bhatia (Mo. #67519)": "",
    "JR Montgomery \t(Mo. #68281)": "",
    "JR Montgomery  (Mo. #68281)": "",
    "Aaron Hadlow\t(Mo. #70987)": "",
    "Aaron Hadlow (Mo. #70987)": "",
    "Hilary Orzick\t(Mo. #78359)": "",
    "Hilary J. Orzick\t(Mo. #78359)": "",
    "Hilary Orzick (Mo. #78359)": "",
    "Hilary J. Orzick (Mo. #78359)": "",
    # Stray blank email lines produced after email substitution
    "": "",
}

# Lines to completely remove (accidentally pasted content, etc.)
REMOVE_PATTERNS = [
    r"^\$env:GOOGLE_",
    r"^node src/server\.js",
    r"^omnibus_notice_of_deposition\s*=",
    r"^corporate_representative_deposition_notice\s*=",
    r"^first_rfps\s*=",
    r"^first_interrogatories\s*=",
    r"^proposed_protective_order\s*=",
    r"^motion_for_entry_of_protective_order\s*=",
]

# File-specific paragraph rules: (prefix_to_match, replacement_text)
# When a paragraph *starts with* the prefix, replace the whole paragraph.
# An empty replacement string removes the paragraph.
FILE_RULES: dict[str, list[tuple[str, str]]] = {
    "rfps-template.docx": [
        # First named-person communication paragraph becomes the block token.
        ("Please produce all communications between {{plaintiffReferenceName}} and Ms. Moore-Jones", "{{rfpPlaintiffCommunicationsBlock}}"),
        ("Please produce all communications between Ms. {{plaintiffReferenceName}} and Ms. Moore-Jones", "{{rfpPlaintiffCommunicationsBlock}}"),
        # Subsequent per-person communication paragraphs are removed (the block handles them).
        ("Please produce all communications between {{plaintiffReferenceName}} and Ms. Jones Scott", ""),
        ("Please produce all communications between {{plaintiffReferenceName}} and Ms. Fockler", ""),
        ("Please produce all communications between {{plaintiffReferenceName}} and Ms. Hafner", ""),
        ("Please produce all communications between {{plaintiffReferenceName}} and Ms. Bradley", ""),
        ("Please produce all communications between {{plaintiffReferenceName}} and Mr. Wise", ""),
        ("Please produce all communications between {{plaintiffReferenceName}} and Ms. Perez", ""),
        ("Please produce all communications between {{plaintiffReferenceName}} and MCC’s Human Resources", ""),
        # Fallback for any remaining raw-name versions
        ("Please produce all communications between Ms. Payne and Ms. Jones Scott", ""),
        ("Please produce all communications between Ms. Payne and Ms. Fockler", ""),
        ("Please produce all communications between Ms. Payne and Ms. Hafner", ""),
        ("Please produce all communications between Ms. Payne and Ms. Bradley", ""),
        ("Please produce all communications between Ms. Payne and Mr. Wise", ""),
        ("Please produce all communications between Ms. Payne and Ms. Perez", ""),
        ("Please produce all communications between Ms. Payne and MCC’s Human Resources", ""),
        # Actor complaint paragraphs
        ("Please produce all communications and documents which relate to any complaints made against or about Ms. Moore-Jones", "{{rfpActorComplaintParagraphsBlock}}"),
        ("Please produce all communications and documents which relate to any complaints made against or about Dr. Jones Scott", ""),
        # Supervisor paragraphs
        ("Please produce all documents that reflect or relate to any coaching, performance counseling", "{{rfpSupervisorCoachingParagraph}}"),
        ("Please produce all documents which constitute lists or charts of all employees supervised or overseen in any way by Ms. Moore-Jones", "{{rfpSupervisorEmployeeListParagraph}}"),
        ("Please produce all documents which reflect the circumstances of the separation from employment of any employee who reported to Ms. Moore-Jones", "{{rfpSupervisorSeparationParagraph}}"),
        # Discipline paragraphs
        ("Please produce all documents and communication related to placing {{plaintiffReferenceName}} on a Formal Verbal Warning", "{{rfpPlaintiffDisciplineParagraphsBlock}}"),
        ("Please produce all documents and communication related to placing Ms. Payne on a Formal Verbal Warning", "{{rfpPlaintiffDisciplineParagraphsBlock}}"),
        ("Please produce all documents and communication related to placing {{plaintiffReferenceName}} on a Performance Improvement Plan", ""),
        ("Please produce all documents and communication related to placing Ms. Payne on a Performance Improvement Plan", ""),
        # Training paragraphs
        ("Please produce all documents reflecting training that Ms. Moore-Jones completed", "{{rfpSupervisorTrainingParagraph}}"),
        ("Please produce all documents reflecting training that any other person with Ms. Moore-Jones’s job title completed", "{{rfpSupervisorPeerTrainingParagraph}}"),
        # TRIO issue paragraphs — first becomes block, rest removed
        ("Produce all annual performance reports, key metrics reports, and any other reports submitted by Metropolitan Community College", "{{rfpTrioIssueParagraphsBlock}}"),
        ("Produce all documents, spreadsheets, databases, enrollment rosters", ""),
        ("Produce all communications between MCC and the United States Department of Education", ""),
        ("Produce all communications, including emails, text messages, instant messages, voicemails, and written notes, between Gabrielle Moore-Jones", ""),
        ("Produce all documents related to MCC’s internal investigation conducted following {{plaintiffReferenceName}}’s August 2023", ""),
        ("Produce all documents related to MCC’s internal investigation conducted following Ms. Payne’s August 2023", ""),
        ("Produce all documents related to MCC’s internal investigation conducted following {{plaintiffReferenceName}}’s August 2024", ""),
        ("Produce all documents related to MCC’s internal investigation conducted following Ms. Payne’s August 2024", ""),
        ("Please produce all documents that reflect the actual funding received from the federal government", ""),
        ("Produce all training materials, policy documents, procedures, guidelines, or instructions provided to TRIO staff", ""),
        ("Produce all documents, communications, or records reflecting any instruction, direction, or request made by Gabrielle Moore-Jones", ""),
        ("Produce all documents reflecting any audit, review, inquiry, correction, amendment, or resubmission", ""),
    ],
    "interrogatories-template.docx": [
        # Actor complaint interrogatory
        ("Please identify all complaints, whether formal or informal, made by any employee against Ms. Moore-Jones", "{{interrogatoryActorComplaintParagraphsBlock}}"),
        # TRIO metrics interrogatory
        ("With respect to the TRIO Program’s annual key metrics reports submitted to the United States Department of Education", "{{interrogatoryTrioMetricsParagraph}}"),
    ],
    "corporate-representative-template.docx": [
        # First topic becomes the block token; subsequent topics are removed.
        ("{{plaintiffName}}’s employment history", "{{corpRepTopicsBlock}}"),
        ("Plaintiff’s employment history", "{{corpRepTopicsBlock}}"),
    ],
    "omnibus-notice-template.docx": [],
}

# Paragraphs whose text (stripped) matches any of these are removed outright.
KNOWN_JUNK_PREFIXES = (
    "$env:GOOGLE_",
    "node src/server.js",
    "omnibus_notice_of_deposition =",
    "corporate_representative_deposition_notice =",
    "first_rfps =",
    "first_interrogatories =",
    "proposed_protective_order =",
    "motion_for_entry_of_protective_order =",
)

# Caption table cell replacements.
CAPTION_CELL = [
    ("FERN PAYNE", "{{plaintiffName}}"),
    (
        "THE METROPOLITAN\nCOMMUNITY COLLEGE\nFOUNDATION et al.",
        "{{defendantName}}",
    ),
    (
        "THE METROPOLITAN\nCOMMUNITY COLLEGE\nFOUNDATION, et al.",
        "{{defendantName}}",
    ),
    ("Case No.:\n2616-CV12184", "Case No.:\n{{caseNumber}}"),
    ("Case No.: 2616-CV12184", "Case No.: {{caseNumber}}"),
    ("Case No.\n2616-CV12184", "Case No.\n{{caseNumber}}"),
    ("Case No. 2616-CV12184", "Case No. {{caseNumber}}"),
    ("2616-CV12184", "{{caseNumber}}"),
]


# -------------------------------------------------------------------
# Helpers
# -------------------------------------------------------------------

def apply_substrings(text: str) -> str:
    for needle, replacement in SUBSTRINGS:
        text = text.replace(needle, replacement)
    return text


def is_junk(text: str) -> bool:
    stripped = text.strip()
    return any(stripped.startswith(p) for p in KNOWN_JUNK_PREFIXES)


def rewrite_paragraph(paragraph, file_rules: list[tuple[str, str]]) -> str | None:
    """
    Apply all replacements to paragraph.text and return the new text.
    Returns None if the paragraph should be removed entirely.
    """
    original = paragraph.text

    # Remove accidentally pasted junk
    if is_junk(original):
        return None

    text = apply_substrings(original)

    # Exact whole-paragraph match (after substring pass)
    stripped = text.strip()
    if stripped in EXACT:
        text = EXACT[stripped]

    # File-specific prefix rules
    for prefix, replacement in file_rules:
        if text.strip().startswith(prefix):
            text = replacement
            break

    return text


def set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace paragraph content with new_text, preserving the first run's formatting."""
    if not paragraph.runs:
        paragraph.text = new_text
        return

    first_run = paragraph.runs[0]
    # Copy formatting from first run before clearing
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)
    first_run.text = new_text


def remove_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def remove_table(table) -> None:
    t = table._element
    t.getparent().remove(t)


def tokenize_paragraphs(paragraphs, file_rules: list[tuple[str, str]]) -> int:
    """Apply all replacements to a list of paragraphs. Returns change count."""
    changes = 0
    # Collect to list first since we may remove elements mid-iteration.
    para_list = list(paragraphs)
    # Track which corp-rep topics have been replaced (only keep the first).
    in_corp_rep_topics = False

    for para in para_list:
        original = para.text
        new_text = rewrite_paragraph(para, file_rules)

        if new_text is None:
            remove_paragraph(para)
            changes += 1
            continue

        # Corp rep topics: after the first topic becomes {{corpRepTopicsBlock}},
        # remove subsequent topic paragraphs (numbered 2–N).
        if new_text == "{{corpRepTopicsBlock}}":
            in_corp_rep_topics = True
            set_paragraph_text(para, new_text)
            changes += 1
            continue

        if in_corp_rep_topics:
            # Stop removing once we hit the signature block or cert of service
            if any(
                original.strip().startswith(s)
                for s in ("Dated:", "CERTIFICATE", "/s/", "KEENAN", "4600", "{{firmAddressBlock}}", "{{attorneyRosterBlock}}")
            ):
                in_corp_rep_topics = False
            else:
                remove_paragraph(para)
                changes += 1
                continue

        if new_text != original:
            set_paragraph_text(para, new_text)
            changes += 1

    return changes


def tokenize_table(table, file_rules: list[tuple[str, str]]) -> int:
    changes = 0
    for row in table.rows:
        for cell in row.cells:
            # Apply caption-level cell replacements
            cell_text = cell.text
            new_cell_text = cell_text
            for needle, replacement in CAPTION_CELL:
                new_cell_text = new_cell_text.replace(needle, replacement)
            if new_cell_text != cell_text:
                # Write the transformed text back into the cell's paragraphs
                cell.paragraphs[0].text = new_cell_text
                for extra in cell.paragraphs[1:]:
                    remove_paragraph(extra)
                changes += 1

            changes += tokenize_paragraphs(cell.paragraphs, file_rules)

            for nested in cell.tables:
                changes += tokenize_table(nested, file_rules)
    return changes


def replace_omnibus_deponent_tables(doc: Document) -> int:
    """
    In the omnibus notice, TABLE 1 onward contain the deponent schedule.
    Replace TABLE 1's content with the block token and delete TABLE 2+.
    TABLE 0 is the caption — leave it alone.
    """
    tables = doc.tables
    if len(tables) <= 1:
        return 0

    changes = 0
    deponent_tables = [t for t in tables[1:] if any("Deponent:" in cell.text for row in t.rows for cell in row.cells)]

    if not deponent_tables:
        return 0

    # Replace the first deponent table's first cell with the block token.
    first = deponent_tables[0]
    first_cell = first.rows[0].cells[0]
    first_cell.paragraphs[0].text = "{{omnibusScheduleBlocks}}"
    for extra_para in first_cell.paragraphs[1:]:
        remove_paragraph(extra_para)
    for extra_row in first.rows[1:]:
        row_el = extra_row._tr
        row_el.getparent().remove(row_el)
    # Clear remaining cells in row 0
    for cell in first.rows[0].cells[1:]:
        for para in cell.paragraphs:
            para.text = ""
    changes += 1

    # Delete tables 2+.
    for extra_table in deponent_tables[1:]:
        remove_table(extra_table)
        changes += 1

    return changes


def tokenize_file(input_path: Path, output_path: Path) -> int:
    output_name = output_path.name
    file_rules = FILE_RULES.get(output_name, [])

    doc = Document(input_path)
    changes = 0

    changes += tokenize_paragraphs(doc.paragraphs, file_rules)

    for table in doc.tables:
        changes += tokenize_table(table, file_rules)

    if output_name == "omnibus-notice-template.docx":
        changes += replace_omnibus_deponent_tables(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return changes


def main() -> None:
    missing = [name for name in FILE_MAP if not (INPUT_DIR / name).exists()]
    if missing:
        print("Missing input files:")
        for name in missing:
            print(f"  {name}")
        return

    for input_name, output_name in FILE_MAP.items():
        input_path = INPUT_DIR / input_name
        output_path = OUTPUT_DIR / output_name
        changes = tokenize_file(input_path, output_path)
        print(f"{output_name}: {changes} replacements")

    print(f"\nTokenized files written to:\n  {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
