import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ALLEGATION_PATTERNS = {
    "Disability discrimination": [r"disability discrimination", r"discriminat(?:ion|e).*disab"],
    "Failure to accommodate": [r"failure to accommodate", r"reasonable accommodation"],
    "Retaliation": [r"\bretaliation\b", r"retaliat(?:e|ion)"],
    "Age discrimination": [r"age discrimination", r"\bage\b.*discriminat"],
    "Race discrimination": [r"race discrimination", r"\brace\b.*discriminat"],
    "Sex discrimination": [r"sex discrimination", r"gender discrimination"],
    "Whistleblower retaliation": [r"whistleblower", r"RSMo\s*105\.055", r"reports? of .*falsification"],
    "Workers' compensation retaliation": [r"workers[’'] compensation", r"workers compensation"],
}

PROTECTED_TRAIT_PATTERNS = {
    "Disability": [r"disability", r"disabled", r"medical condition", r"accommodation"],
    "Age": [r"\bage\b", r"older", r"younger"],
    "Race": [r"\brace\b"],
    "Sex": [r"\bsex\b", r"\bgender\b"],
}

RETALIATION_ACTIVITY_PATTERNS = {
    "Requests for accommodation": [r"request(?:s)? for reasonable accommodations?", r"reasonable accommodation"],
    "Protected leave requests": [r"FMLA", r"protected leave", r"medical leave"],
    "Reports of discrimination": [r"complaints? of discrimination", r"reported discrimination"],
    "Reports of retaliation": [r"complaints? of retaliation", r"reported retaliation"],
    "Reports of data falsification": [r"data falsification", r"falsification of .*data", r"enrollment data"],
    "Whistleblower reports": [r"whistleblower", r"RSMo\s*105\.055", r"suspected fraud"],
}

ADVERSE_ACTION_PATTERNS = {
    "Discipline": [r"\bdiscipline\b", r"formal verbal warning"],
    "Performance Improvement Plan": [r"Performance Improvement Plan", r"\bPIP\b"],
    "Schedule reduction": [r"reduction of .*work schedule", r"five days per week to two days per week", r"compensation reduction"],
    "Termination": [r"\btermination\b", r"terminated", r"separation from employment"],
    "Denial of accommodation": [r"denial of .*accommodation", r"denied.*accommodation"],
}

COMPLAINT_TYPE_PATTERNS = {
    "Retaliation complaints": [r"complaints? or reports? of retaliation"],
    "Disability discrimination complaints": [r"complaints? or reports? of disability discrimination"],
    "Age discrimination complaints": [r"complaints? or reports? of age discrimination"],
    "Whistleblower complaints": [r"whistleblower", r"violation of whistleblowers[’'] rights", r"RSMo\s*105\.055"],
    "Workers' compensation retaliation complaints": [r"workers[’'] compensation-related retaliation"],
}

ROLE_PATTERNS = {
    "Persons involved in discipline decisions": [r"decision to discipline", r"discipline.*decision"],
    "Persons involved in accommodation decisions": [r"accommodation process", r"requests? for reasonable accommodations?"],
    "TRIO leadership": [r"TRIO leadership", r"TRIO program"],
    "Human resources personnel": [r"Human Resources", r"\bHR\b"],
}

COMPARATOR_PATTERNS = {
    "Employees in the relevant business unit": [r"TRIO program", r"business unit"],
    "Employees placed on performance improvement plans": [r"placed on a Performance Improvement Plan", r"\bPIP\b"],
    "Employees granted remote-work accommodations": [r"granted remote-work accommodations", r"remote work accommodation"],
}

NAME_STOPWORDS = {
    "Plaintiff",
    "First",
    "Set",
    "Defendant",
    "Defendants",
    "Court",
    "Circuit",
    "Complaint",
    "Petition",
    "Count",
    "Counts",
    "Charge",
    "Discrimination",
    "Retaliation",
    "Missouri",
    "Kansas",
    "City",
    "Jackson",
    "United States",
    "Performance",
    "Improvement",
    "Plan",
    "Case",
    "No",
    "Attorney",
    "Attorneys",
    "Madison",
    "Ave",
    "College",
    "Community",
    "Metropolitan",
    "Foundation",
    "Human",
    "Resources",
    "TRIO",
}


def read_pdf_text(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def read_docx_text(file_path: Path) -> str:
    document = Document(str(file_path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_text)


def read_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf_text(file_path)
    if suffix == ".docx":
        return read_docx_text(file_path)
    return file_path.read_text(encoding="utf8", errors="ignore")


def normalized(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def detect_values(text: str, pattern_map: dict[str, list[str]]) -> list[str]:
    text_lower = text.lower()
    found = []
    for label, patterns in pattern_map.items():
        if any(re.search(pattern, text_lower, flags=re.IGNORECASE) for pattern in patterns):
            found.append(label)
    return found


def extract_case_number(text: str) -> str:
    match = re.search(r"Case No\.?:?\s*([A-Za-z0-9\-]+)", text, flags=re.IGNORECASE)
    return match.group(1) if match else ""


def extract_counts(text: str) -> list[str]:
    counts = []
    for match in re.finditer(r"(Count\s+[IVXLC0-9]+[:.\-\s]+)(.{0,180})", text, flags=re.IGNORECASE):
        label = normalized(f"{match.group(1)}{match.group(2)}")
        counts.append(label)
    return counts


def extract_people(text: str) -> list[str]:
    matches = re.findall(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})\b", text)
    cleaned = []
    for name in matches:
        parts = [part.strip() for part in name.split()]
        if len(parts) < 2:
            continue
        if any(part in NAME_STOPWORDS for part in parts):
            continue
        if any(len(part) == 1 for part in parts):
            continue
        cleaned.append(" ".join(parts))
    counts = Counter(cleaned)
    people = [name for name, count in counts.most_common(25) if count >= 2]
    return people[:12]


def extract_dates(text: str) -> list[str]:
    patterns = [
        r"(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}",
        r"\b\d{1,2}/\d{1,2}/\d{2,4}\b",
        r"\b20\d{2}\b",
    ]
    found = []
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            value = normalized(match.group(0))
            if value not in found:
                found.append(value)
    return found[:15]


def derive_corp_rep_topics(suggestions: dict, people: list[str]) -> list[str]:
    topics = []
    if suggestions["claimsAndCounts"]:
        topics.append(f"All facts relating to the claims and counts asserted in this matter, including {', '.join(suggestions['claimsAndCounts'][:4])}.")
    if suggestions["retaliationActivities"]:
        topics.append(f"All facts relating to Plaintiff's protected activities, including {', '.join(suggestions['retaliationActivities'][:4])}.")
    if suggestions["adverseActions"]:
        topics.append(f"All facts relating to the adverse actions at issue, including {', '.join(suggestions['adverseActions'][:4])}.")
    if suggestions["complaintTypes"]:
        topics.append(f"All complaints or reports involving {', '.join(suggestions['complaintTypes'][:5])}.")
    if people:
        topics.append(f"The identities and roles of the following persons and any other individuals involved in the challenged events: {', '.join(people[:8])}.")
    return topics[:8]


def derive_interrogatory_prompts(suggestions: dict, people: list[str]) -> list[str]:
    prompts = []
    if people:
        prompts.append(f"Complaints, disciplinary actions, or allegations involving the identified persons, including {', '.join(people[:6])}.")
    if suggestions["complaintTypes"]:
        prompts.append(f"Individuals who reported or were associated with {', '.join(suggestions['complaintTypes'][:5])}.")
    if suggestions["decisionMakers"]:
        prompts.append(f"Persons who participated in {', '.join(suggestions['decisionMakers'][:4])}.")
    if suggestions["comparatorGroups"]:
        prompts.append(f"Comparator employees or groups, including {', '.join(suggestions['comparatorGroups'][:4])}.")
    return prompts[:6]


def build_summary(suggestions: dict, counts: list[str], people: list[str], case_number: str, docs: list[dict]) -> list[str]:
    summary = []
    if case_number:
        summary.append(f"Detected case number: {case_number}")
    if counts:
        summary.append(f"Detected counts: {', '.join(counts[:6])}")
    if suggestions["claimsAndCounts"]:
        summary.append(f"Likely claims: {', '.join(suggestions['claimsAndCounts'])}")
    if people:
        summary.append(f"Likely deponents / relevant actors: {', '.join(people[:8])}")
    summary.append(f"Parsed {len(docs)} uploaded document(s). Review all extracted values before using them in filings.")
    return summary


def main():
    payload = json.load(sys.stdin)
    files = payload.get("files", [])

    documents = []
    combined_text_parts = []

    for item in files:
        file_path = Path(item["path"])
        text = read_text(file_path)
        trimmed = normalized(text)
        documents.append(
            {
                "name": item["name"],
                "path": str(file_path),
                "characters": len(trimmed),
                "preview": trimmed[:500],
            }
        )
        combined_text_parts.append(trimmed)

    combined_text = "\n".join(combined_text_parts)
    case_number = extract_case_number(combined_text)
    counts = extract_counts(combined_text)
    people = extract_people(combined_text)

    suggestions = {
        "caseNumber": case_number,
        "claimsAndCounts": counts or detect_values(combined_text, ALLEGATION_PATTERNS),
        "protectedTraits": detect_values(combined_text, PROTECTED_TRAIT_PATTERNS),
        "retaliationActivities": detect_values(combined_text, RETALIATION_ACTIVITY_PATTERNS),
        "adverseActions": detect_values(combined_text, ADVERSE_ACTION_PATTERNS),
        "complaintTypes": detect_values(combined_text, COMPLAINT_TYPE_PATTERNS),
        "decisionMakers": detect_values(combined_text, ROLE_PATTERNS),
        "comparatorGroups": detect_values(combined_text, COMPARATOR_PATTERNS),
        "keyPersonsList": people,
        "corpRepIssueTopics": [],
        "interrogatoryIssuePrompts": [],
        "specialInstructions": "",
    }

    suggestions["corpRepIssueTopics"] = derive_corp_rep_topics(suggestions, people)
    suggestions["interrogatoryIssuePrompts"] = derive_interrogatory_prompts(suggestions, people)

    if people:
        suggestions["specialInstructions"] = (
            f"Review extracted actor list and deposition targets: {', '.join(people[:8])}."
        )

    response = {
        "ok": True,
        "summary": build_summary(suggestions, counts, people, case_number, documents),
        "suggestions": {key: "\n".join(value) if isinstance(value, list) else value for key, value in suggestions.items() if value},
        "documents": documents,
        "detectedDates": extract_dates(combined_text),
    }

    json.dump(response, sys.stdout, indent=2)


if __name__ == "__main__":
    main()
