from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Users\noah\Documents\GitHub\first-day-filings")
SOURCE_DIR = Path(r"C:\Users\noah\Downloads")
OUTPUT_DIR = ROOT / "templates" / "docx-masters"


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def remove_table(table):
    table._element.getparent().remove(table._element)


def replace_in_paragraph(paragraph, replacements):
    text = paragraph.text
    for old, new in replacements.items():
        text = text.replace(old, new)
    paragraph.text = text


def replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)


def paragraph_index(doc, needle):
    for index, paragraph in enumerate(doc.paragraphs):
        if needle in paragraph.text:
            return index
    raise ValueError(f"Paragraph containing {needle!r} not found")


def paragraph_by_text(doc, needle):
    return doc.paragraphs[paragraph_index(doc, needle)]


def paragraph_index_startswith(doc, needle):
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(needle):
            return index
    raise ValueError(f"Paragraph starting with {needle!r} not found")


def paragraph_by_startswith(doc, needle):
    return doc.paragraphs[paragraph_index_startswith(doc, needle)]


def find_paragraph_containing(doc, needle):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    return None


def replace_block_with_placeholder(doc, first_paragraph_text, last_paragraph_text, placeholder):
    start = paragraph_index(doc, first_paragraph_text)
    end = paragraph_index(doc, last_paragraph_text)
    doc.paragraphs[start].text = placeholder
    for idx in range(end, start, -1):
        remove_paragraph(doc.paragraphs[idx])


def replace_signature_block(doc, attorney_for_text):
    dated_idx = paragraph_index(doc, "Dated:")
    firm_idx = paragraph_index(doc, "KEENAN & BHATIA")
    sign_idx = paragraph_index_startswith(doc, "/s/")
    roster_start = paragraph_index(doc, "Edward (E.E.) Keenan")
    roster_end = paragraph_index(doc, "Hilary")
    address_start = paragraph_index(doc, "4600 Madison Ave.")
    address_end = paragraph_index(doc, "(816) 809-2100")
    email_start = paragraph_index(doc, "ee@keenanfirm.com")
    email_end = paragraph_index(doc, "hilary@keenanfirm.com") if any("hilary@keenanfirm.com" in p.text for p in doc.paragraphs) else paragraph_index(doc, "aaron@keenanfirm.com")
    attorney_for_idx = paragraph_index(doc, attorney_for_text)

    doc.paragraphs[dated_idx].text = "Dated: {{serviceDate}}\t\t\tRespectfully submitted,"
    doc.paragraphs[firm_idx].text = "{{firmName}}"
    doc.paragraphs[sign_idx].text = "/s/{{signingAttorney}}"
    doc.paragraphs[roster_start].text = "{{attorneyRosterBlock}}"
    for idx in range(roster_end, roster_start, -1):
        remove_paragraph(doc.paragraphs[idx])
    doc.paragraphs[address_start].text = "{{firmAddressBlock}}"
    for idx in range(address_end, address_start, -1):
        remove_paragraph(doc.paragraphs[idx])
    doc.paragraphs[email_start].text = "{{attorneyEmailBlock}}"
    for idx in range(email_end, email_start, -1):
        remove_paragraph(doc.paragraphs[idx])
    doc.paragraphs[attorney_for_idx].text = "{{attorneyForLine}}"


def replace_caption(doc):
    replacements = {
        "IN THE CIRCUIT COURT OF JACKSON COUNTY, MISSOURI": "{{courtName}}",
        "IN THE CIRCUIT COURT OF JACKSON COUNTY": "{{courtName}}",
        "AT KANSAS CITY": "{{courtDivision}}",
        "FERN PAYNE": "{{plaintiffName}}",
        "THE METROPOLITAN COMMUNITY COLLEGE FOUNDATION et al.": "{{defendantName}}",
        "THE METROPOLITAN COMMUNITY COLLEGE FOUNDATION, et al.": "{{defendantName}}",
        "Case No.: 2616-CV12184": "Case No.: {{caseNumber}}",
        "Case No. 2616-CV12184": "Case No. {{caseNumber}}",
    }
    for paragraph in doc.paragraphs[:3]:
        replace_in_paragraph(paragraph, replacements)
    replace_in_table(doc.tables[0], replacements)


def create_omnibus_template():
    doc = Document(SOURCE_DIR / "Plaintiff Last Name_Defendant - Omnibus Notice of Deposition.docx")
    replace_caption(doc)
    replace_in_paragraph(
        paragraph_by_text(doc, "Please take notice that Plaintiff Fern Payne"),
        {"Fern Payne": "{{plaintiffName}}"},
    )
    intro_paragraph = paragraph_by_text(doc, "Please take notice that Plaintiff")
    insert_paragraph_after(intro_paragraph, "{{omnibusScheduleBlocks}}", intro_paragraph.style)
    for table in list(doc.tables[1:]):
        remove_table(table)
    replace_in_paragraph(paragraph_by_text(doc, "These depositions will be taken stenographically"), {"Noah Tunis": "{{videoOperatorName}}"})
    replace_in_paragraph(paragraph_by_text(doc, "I certify that on"), {"April 20, 2026": "{{serviceDate}}"})
    paragraph_by_text(doc, "/s/ Aaron Hadlow").text = "/s/{{signingAttorney}}"
    paragraph_by_text(doc, "An Attorney for Plaintiff").text = "An Attorney for Plaintiff {{plaintiffName}}"
    replace_signature_block(doc, "Attorneys for Plaintiff Fern Payne")
    output = OUTPUT_DIR / "omnibus-notice-template.docx"
    doc.save(output)


def create_corp_rep_template():
    doc = Document(SOURCE_DIR / "Plaintiff Last Name _ Defendant - Notice of Corporate Representative Deposition.docx")
    replace_caption(doc)
    paragraph_by_text(doc, "Deponent:").text = "Deponent:\t{{corpRepEntity}}"
    paragraph_by_text(doc, "Location:").text = "Location:\t{{corpRepLocation}}"
    paragraph_by_text(doc, "Date/Time:").text = "Date/Time:  {{corpRepDateTime}}"
    paragraph_by_text(doc, "Format:").text = "Format:  {{corpRepFormat}}"
    paragraph_by_text(doc, "Pursuant to Mo. S. Ct R. 57.03(b)(3)").text = "{{corpRepDocumentRequest}}"
    replace_in_paragraph(paragraph_by_text(doc, "I certify that on"), {"April 20, 2026": "{{serviceDate}}"})
    attorney_line = find_paragraph_containing(doc, "Attorney for Plaintiff")
    if attorney_line is not None:
        attorney_line.text = "Attorney for Plaintiff {{plaintiffName}}"
    replace_block_with_placeholder(
        doc,
        "Plaintiff’s employment history, job duties, performance evaluations, and disciplinary history from 2003 to present.",
        "The demographic makeup of the MCC location where Plaintiff worked, including the age, race, color, sex, and disability status of its employees.",
        "{{corpRepTopicsBlock}}",
    )
    replace_signature_block(doc, "Attorneys for Plaintiff Fern Payne")
    output = OUTPUT_DIR / "corporate-representative-template.docx"
    doc.save(output)


def create_rog_template():
    doc = Document(SOURCE_DIR / "Plaintiff Last Name_Defendant - P's 1st ROGs to MCC.docx")
    replace_caption(doc)
    paragraph_by_text(doc, "(First Set)").text = "{{interrogatorySetLabel}}"
    paragraph_by_text(doc, "Plaintiff Fern Payne respectfully directs").text = (
        "Plaintiff {{plaintiffName}} respectfully directs the following {{interrogatorySetLabel}} "
        "Interrogatories to Defendants {{targetDefendants}} pursuant to Mo. S. Ct. R. 57.01. "
        "{{responseDeliveryInstructions}}"
    )
    replace_block_with_placeholder(
        doc,
        "Fern Payne;",
        "Any person who replaced Ms. Payne or assumed any part of her job functions following her departure from MCC.",
        "{{interrogatory3SubjectsBlock}}",
    )
    replace_signature_block(doc, "Attorneys for Plaintiff Fern Payne")
    attorney_line = find_paragraph_containing(doc, "Attorney for Plaintiff")
    if attorney_line is not None:
        attorney_line.text = "Attorney for Plaintiff {{plaintiffName}}"
    paragraph_by_text(doc, "STATE OF ________________").text = "STATE OF {{verificationState}}\t)"
    paragraph_by_text(doc, "COUNTY OF ______________").text = "COUNTY OF {{verificationCounty}}\t)"
    paragraph_by_text(doc, "Now on this").text = (
        "Now on this {{verificationDay}} day of {{verificationMonth}}, {{verificationYear}}, "
        "comes {{verificationAffiantName}}, and verifies under oath that the answers to the above "
        "Interrogatories upon behalf of Defendant {{verificationEntity}}, are true, complete, and correct."
    )
    paragraph_by_text(doc, "Signature of Verifying Person").text = "{{verificationAffiantName}}\t\t{{verificationCapacity}}"
    paragraph_by_text(doc, "Notary Public Signature").text = "Notary Public Signature"
    paragraph_by_text(doc, "Notary Commission Expiration Date").text = "{{notaryExpirationDate}}"
    output = OUTPUT_DIR / "interrogatories-template.docx"
    doc.save(output)


def create_rfp_template():
    doc = Document(SOURCE_DIR / "Plaintiff Last Name_Defendant - P's 1st RFPs to MCC.docx")
    replace_caption(doc)
    paragraph_by_text(doc, "(First Set)").text = "{{rfpSetLabel}}"
    paragraph_by_text(doc, "Plaintiff Fern Payne directs the following Requests").text = (
        "Plaintiff {{plaintiffName}} directs the following Requests for Production of Documents and Things "
        "to Defendants {{targetDefendants}}. {{productionFormatInstructions}}"
    )
    paragraph_by_text(doc, "References to  Defendants The Metropolitan Community College Foundation").text = (
        "References to Defendants {{targetDefendants}} include any parents, subsidiaries, and affiliates. "
        "As used here, “{{collectiveDefendantShortName}}” and “Defendants” refer to any of the Defendant entities "
        "and any of their parents, subsidiaries, and affiliates. If Plaintiff is asking about a smaller business unit, "
        "that is specified."
    )
    paragraph_by_text(doc, "This document refers to Plaintiff as Fern Payne.").text = (
        "This document refers to Plaintiff as {{plaintiffName}}. However, requests here refer to Plaintiff, "
        "regardless of what name or nickname she was called."
    )
    paragraph_by_text(doc, "Electronically stored information").text = "{{esiInstructionText}}"
    replace_block_with_placeholder(
        doc,
        "Fern Payne;",
        "Any person whom you may call to testify or otherwise give evidence (e.g., a statement or affidavit) in this action.",
        "{{rfpSubjectListBlock}}",
    )
    replace_block_with_placeholder(
        doc,
        "Fern Payne;",
        "Any person whom you may call to testify or otherwise give evidence (e.g., a statement or affidavit) in this action.",
        "{{rfpSubjectListBlock}}",
    )
    replace_block_with_placeholder(
        doc,
        "Fern Payne;",
        "Any person whom you may call to testify or otherwise give evidence (e.g., a statement or affidavit) in this action.",
        "{{rfpSubjectListBlock}}",
    )
    replace_block_with_placeholder(
        doc,
        "Fern Payne;",
        "Any person who replaced Ms. Payne or assumed any part of her job functions following her departure from MCC.",
        "{{rfpEmailBoxSubjectListBlock}}",
    )
    replace_signature_block(doc, "Attorneys for Fern Payne")
    paragraph_by_text(doc, "Attorney for Plaintiff").text = "Attorney for Plaintiff {{plaintiffName}}"
    output = OUTPUT_DIR / "rfps-template.docx"
    doc.save(output)


def create_protective_order_template():
    doc = Document(SOURCE_DIR / "Plaintiff Last Name_Defendant - Proposed Protective Order.docx")
    replace_caption(doc)
    replace_in_paragraph(
        paragraph_by_text(doc, "The allegations and defenses in this case"),
        {
            "Plaintiff Fern Payne": "Plaintiff {{plaintiffName}}",
            "The Metropolitan Community College Foundation, Metropolitan Community College, and The Junior College District of Metropolitan Kansas City": "{{targetDefendants}}",
            "“MCC”": "“{{collectiveDefendantShortName}}”",
        },
    )
    paragraph_by_text(doc, "a.  “Action” means the above-entitled proceedings").text = (
        "a.  “Action” means the above-entitled proceedings: {{protectiveOrderActionName}} and any appeals therefrom."
    )
    paragraph_by_text(doc, "c.\xa0  “Confidential Materials” means").text = (
        "c.  “Confidential Materials” means {{confidentialMaterialsDefinition}}"
    )
    replace_in_paragraph(paragraph_by_text(doc, "i.  “Party” means Plaintiff Fern Payne"), {"Plaintiff Fern Payne": "Plaintiff {{plaintiffName}}"})
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, {"HON. ADAM CAINE": "{{judgeName}}"})
    output = OUTPUT_DIR / "protective-order-template.docx"
    doc.save(output)


def create_motion_for_protective_order_template():
    doc = Document(SOURCE_DIR / "Payne_MCC - Motion for Entry of Protective Order.docx")
    replace_caption(doc)
    replace_in_paragraph(
        paragraph_by_text(doc, "Plaintiff Fern Payne hereby moves"),
        {"Plaintiff Fern Payne": "Plaintiff {{plaintiffName}}"},
    )
    paragraph_by_text(doc, "Dated:").text = "Dated: {{serviceDate}}\t\t\tRespectfully submitted,"
    paragraph_by_text(doc, "KEENAN & BHATIA, LLC").text = "{{firmName}}"
    paragraph_by_startswith(doc, "/s/").text = "/s/{{signingAttorney}}"
    paragraph_by_text(doc, "4600 Madison Ave.").text = "{{firmAddressBlock}}"
    for needle in ["Kansas City, Missouri 64112", "Tel: (816) 809-2100"]:
        remove_paragraph(paragraph_by_text(doc, needle))
    paragraph_by_text(doc, "sonal@keenanfirm.com").text = "{{attorneyEmailBlock}}"
    for needle in ["ee@keenanfirm.com", "jr@keenanfirm.com", "aaron@keenanfirm.com", "hilary@keenanfirm.com"]:
        remove_paragraph(paragraph_by_text(doc, needle))
    paragraph_by_text(doc, "Attorneys for Plaintiff Fern Payne").text = "{{attorneyForLine}}"
    replace_in_paragraph(paragraph_by_text(doc, "I certify that on"), {"April 20, 2026": "{{serviceDate}}"})
    paragraph_by_text(doc, "/s/ Aaron Hadlow").text = "/s/{{signingAttorney}}"
    paragraph_by_text(doc, "Attorney for Plaintiff Fern Payne").text = "Attorney for Plaintiff {{plaintiffName}}"

    roster_table = doc.tables[1]
    first_cell = roster_table.cell(0, 0)
    first_cell.text = "{{attorneyRosterBlock}}"
    for row_idx in range(len(roster_table.rows) - 1, 0, -1):
        roster_table._tbl.remove(roster_table.rows[row_idx]._tr)
    roster_table.cell(0, 1).text = ""

    output = OUTPUT_DIR / "motion-for-protective-order-template.docx"
    doc.save(output)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    create_omnibus_template()
    create_corp_rep_template()
    create_rog_template()
    create_rfp_template()
    create_protective_order_template()
    create_motion_for_protective_order_template()
    print(f"Wrote templates to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
